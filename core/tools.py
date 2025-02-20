import os
import subprocess
import platform
import re
import uuid
import PyPDF2
import docx
from concurrent.futures import ThreadPoolExecutor
from multiprocessing import Pool, Manager, cpu_count
from langchain_core.tools import StructuredTool, tool
from langchain_core.documents import Document
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_openai import OpenAIEmbeddings
from dotenv import load_dotenv
try:
    from pydantic.v1 import BaseModel, Field
except ImportError:
    from pydantic import BaseModel, Field

load_dotenv()

api_key = os.getenv("OPENAI_API_KEY")

if not api_key:
    raise ValueError("Error: OPENAI_API_KEY is not set in .env!")

embeddings = OpenAIEmbeddings(
    openai_api_key=api_key, model="text-embedding-3-large"
)

EXCLUDED_FOLDERS = {
    "node_modules", ".git", ".venv", "venv", "__pycache__", ".DS_Store",
    "Library", "System", "Applications", "usr", "bin", "opt", "var", ".Trash"
}


class ToolArgs(BaseModel):
    """Schema to ensure argument validation works correctly."""
    param: str = Field(..., description="Input parameter")


def create_tool(name: str, description: str, func):
    """Dynamically creates a StructuredTool, ensuring compatibility with Pydantic v1 & v2."""
    formatted_name = re.sub(r"[^\w\s]", "", name).replace(" ", "_")

    return StructuredTool.from_function(
        func, name=formatted_name, description=description, args_schema=ToolArgs
    )


TOOL_REGISTRY = {}


def register_tool(name: str, description: str, func):
    """Registers a tool dynamically with an embedding."""
    tool_id = str(uuid.uuid4())
    TOOL_REGISTRY[tool_id] = create_tool(name, description, func)
    return tool_id


def get_tool_registry():
    """Returns the tool registry."""
    return TOOL_REGISTRY


def initialize_vector_store():
    """Initializes a vector store with tool descriptions."""
    tool_documents = [
        Document(page_content=tool.description, id=tool_id,
                 metadata={"tool_name": tool.name})
        for tool_id, tool in TOOL_REGISTRY.items()
    ]

    vector_store = InMemoryVectorStore(embedding=embeddings)
    vector_store.add_documents(tool_documents)

    return vector_store


VECTOR_STORE = None


@tool
def resolve_path(path: str) -> str:
    """Resolves a user-provided path to its absolute system path."""
    if path.lower() == "desktop":
        return os.path.join(os.path.expanduser("~"), "Desktop")
    elif path.lower() == "documents":
        return os.path.join(os.path.expanduser("~"), "Documents")
    return os.path.expanduser(path)


register_tool("Resolve Path",
              "Resolves relative paths to absolute paths.", resolve_path)


@tool
def create_folder(folder_name: str, path: str = ".") -> str:
    """Creates a new folder at the specified path."""
    full_path = os.path.join(resolve_path(path), folder_name)
    try:
        os.makedirs(full_path, exist_ok=True)
        return f"Folder '{folder_name}' created at {full_path}."
    except Exception as e:
        return f"Error creating folder '{folder_name}': {e}"


register_tool("Create Folder",
              "Creates a folder at a specified location.", create_folder)


@tool
def create_file(file_name: str, path: str = ".") -> str:
    """Creates an empty file at the specified path."""
    full_path = os.path.join(resolve_path(path), file_name)
    try:
        with open(full_path, 'w') as file:
            file.write("")
        return f"File '{file_name}' created at {full_path}."
    except Exception as e:
        return f"Error creating file '{file_name}': {e}"


register_tool("Create File",
              "Creates a new file at a given path.", create_file)


@tool
def list_files_and_folders(path: str = ".") -> list:
    """Lists all files and folders in a specified directory."""
    try:
        return os.listdir(resolve_path(path))
    except Exception as e:
        return [f"Error accessing the directory: {e}"]


register_tool("List Files and Folders",
              "Lists all files and directories in a path.", list_files_and_folders)


@tool
def read_file_content(file_path: str) -> str:
    """Reads the content of `.txt`, `.pdf`, and `.docx` files."""
    file_extension = os.path.splitext(file_path)[1].lower()
    content = ""

    try:
        if file_extension == ".txt":
            with open(file_path, 'r') as file:
                content = file.read()
        elif file_extension == ".pdf":
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    content += page.extract_text() or ""
        elif file_extension == ".docx":
            doc = docx.Document(file_path)
            content = "\n".join([para.text for para in doc.paragraphs])
        else:
            return f"Unsupported file format: {file_extension}"
    except Exception as e:
        return f"Failed to read file: {e}"

    return content


register_tool("Read File Content",
              "Reads content from .txt, .pdf, and .docx files.", read_file_content)


@tool
def append_to_file(file_path: str, content: str) -> str:
    """Appends text to an existing file."""
    resolved_path = resolve_path(file_path)

    if not os.path.exists(resolved_path):
        return f"Error: File '{resolved_path}' does not exist."

    try:
        with open(resolved_path, 'a') as file:
            file.write(content + "\n")
        return f"Content appended to {resolved_path} successfully."
    except Exception as e:
        return f"Error appending to file '{resolved_path}': {e}"


register_tool("Append to File",
              "Appends content to a specified file.", append_to_file)


@tool
def search_for_target(target_name: str, search_path: str = None) -> list:
    """Searches for a file or folder in parallel across all directories."""
    if search_path is None:
        search_path = os.path.expanduser("~")

    max_workers = cpu_count()
    manager = Manager()
    queue = manager.Queue()
    queue.put(search_path)
    found_paths = manager.list()

    with Pool(processes=max_workers) as pool:
        while not queue.empty():
            tasks = []
            for _ in range(min(queue.qsize(), max_workers)):
                current_dir = queue.get()
                tasks.append(pool.apply_async(
                    search_directory, (current_dir, target_name)))

            for task in tasks:
                result, new_dirs = task.get()
                if result:
                    found_paths.extend(result)
                for new_dir in new_dirs:
                    if should_exclude(new_dir):
                        continue
                    queue.put(new_dir)

    return list(found_paths) if found_paths else []


register_tool("Search for Target",
              "Searches for a file or folder in a given directory.", search_for_target)


@tool
def search_and_append_to_file(file_name: str, content: str, search_path: str = None) -> str:
    """Searches for a file and appends content to it if found."""
    if search_path is None:
        search_path = os.path.expanduser("~")  # Default to home directory

    found_files = search_for_target(file_name, search_path, is_file=True)

    if len(found_files) == 1:
        return append_to_file(found_files[0], content)
    elif len(found_files) > 1:
        return f"Multiple files found:\n" + "\n".join(found_files)
    else:
        return f"File '{file_name}' not found in the search path."


register_tool("Search and Append to File",
              "Searches for a specific file and appends content to it.", search_and_append_to_file)


@tool
def open_file_or_folder(target_name: str, search_path: str = None) -> str:
    """Searches for and opens a file or folder if found."""
    if search_path is None:
        search_path = os.path.expanduser("~")

    found_targets = search_for_target(target_name, search_path)

    if not found_targets:
        return f"'{target_name}' not found."

    if len(found_targets) > 1:
        return f"Multiple matches found:\n" + "\n".join(found_targets)

    target_path = found_targets[0]

    try:
        system = platform.system()
        if system == "Windows":
            subprocess.run(["start", target_path], shell=True)
        elif system == "Darwin":
            subprocess.run(["open", target_path])
        elif system == "Linux":
            subprocess.run(["xdg-open", target_path])
        else:
            return "Unsupported operating system."
        return f" '{target_name}' opened successfully: {target_path}"
    except Exception as e:
        return f" Error opening '{target_name}': {e}"


register_tool("Open File or Folder",
              "Opens a file or folder in the system.", open_file_or_folder)


### HELPER FUNCTIONS ###
def should_exclude(path):
    """Checks if a directory should be skipped to speed up search."""
    return any(excluded in path.split(os.sep) for excluded in EXCLUDED_FOLDERS)


def search_directory(path, target_name):
    """Scans a directory for a file or folder and returns new directories to search."""
    if should_exclude(path):
        return [], []

    found_paths = []
    subdirectories = []

    try:
        with os.scandir(path) as entries:
            for entry in entries:
                if entry.name == target_name:
                    found_paths.append(entry.path)
                elif entry.is_dir(follow_symlinks=False):
                    subdirectories.append(entry.path)

    except (PermissionError, OSError):
        pass

    return found_paths, subdirectories
